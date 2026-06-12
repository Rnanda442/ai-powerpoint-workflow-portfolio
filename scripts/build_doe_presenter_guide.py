from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "deliverables" / "DOE_AI_Portfolio_Presenter_Guide_v2.docx"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "172033"
MUTED = "64748B"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_YELLOW = "FFF7ED"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_run(run, bold=False, italic=False, color=INK, size=None):
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.style = "Normal"
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        style_run(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        style_run(r2)
    else:
        r = p.add_run(text)
        style_run(r)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    style_run(r)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    style_run(r)
    return p


def add_callout(doc, title, text, fill=PALE_YELLOW):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    style_run(r, bold=True, color=DARK_BLUE)
    p.add_run("\n")
    r2 = p.add_run(text)
    style_run(r2)
    doc.add_paragraph()


TOPICS = [
    {
        "slug": "ai_workflow",
        "title": "How AI Agents Learn Scientific Software",
        "quick": "Screenshots, task recordings, file state, and rubrics can become training examples for agents that use QGIS, ParaView, and other scientific tools.",
        "deep": "This topic is not just about using ChatGPT. It is about turning human scientific software work into structured demonstrations. I can show the task, the screen state, what I clicked, what output counted as correct, and where human review still matters.",
        "doe": "DOE relevance: national labs and contractors use specialized tools. The bottleneck is not only computation; it is getting agents to operate real scientific software safely and measurably.",
        "models": [
            "OCR / UI state encoder: reads visible text, buttons, map layers, file trees, and errors so the screen becomes model-readable.",
            "Behavior-cloning transformer: learns action sequences from demonstrations, similar to observe screen -> act -> inspect result.",
            "Replay simulator: reruns the task so success is judged by behavior, not by a pretty final screenshot.",
        ],
        "questions": [
            ("How do you know the agent actually learned the task?", "I would hold out tasks, files, and projects, then score replay success against a rubric. Final screenshots are not enough."),
            ("What should be automated first?", "Repetitive, low-risk setup and QA steps: opening layers, checking CRS, exporting maps, running known workflows, and flagging failures."),
            ("What remains human?", "Scientific acceptance criteria, unsafe file actions, ambiguous outputs, and any interpretation that changes a decision."),
        ],
        "sectors": [
            "Finance: analyst workflow agents that gather filings and produce auditable research packets.",
            "Marketing: campaign-ops agents that reproduce dashboard/report workflows with approval gates.",
            "AI startups: vertical agents for niche software, trained on expert demonstrations.",
            "Jobs: workflow evaluator, rubric designer, agent QA analyst.",
        ],
        "fresh": "The new angle is that training data is not just text. It is the workflow itself: screen, files, clicks, rubric, and reviewed output.",
    },
    {
        "slug": "thesis_graph",
        "title": "AI + Knowledge Graphs For Critical Minerals",
        "quick": "Critical-mineral research can move from scattered papers and slides into a graph of deposits, minerals, host rocks, processes, and evidence.",
        "deep": "Here I would explain Mountain Pass and Bayan Obo as a source-to-graph workflow. The useful part is not claiming AI discovers a deposit. The useful part is making relationships visible and reviewable.",
        "doe": "DOE relevance: critical minerals affect energy storage, defense supply chains, magnets, and clean-energy technologies. Graph structure can help experts find patterns faster while preserving evidence.",
        "models": [
            "SciBERT / MatSciBERT: scientific language models that tag mineral names, host rocks, processes, and evidence phrases.",
            "Relation cross-encoder: looks at two entities and a source sentence, then ranks whether the relationship is supported.",
            "GraphRAG: answers questions by following graph links back to cited evidence.",
            "GraphSAGE / R-GCN: graph neural networks that update nodes or edges using nearby relationships.",
        ],
        "questions": [
            ("Where does a graph beat a spreadsheet?", "When relationships matter: mineral hosted by rock, altered by fluid, supported by source, similar to another deposit."),
            ("Can graph ML discover deposits?", "I would not claim that. It can rank relationships or hypotheses for expert review."),
            ("What would you validate?", "Entity cleanup, edge type, evidence weight, source citation, and whether inferred links are geologically meaningful."),
        ],
        "sectors": [
            "Finance: company, supplier, customer, risk, and filing relationship graphs.",
            "Marketing: customer journey graphs connecting touchpoints, products, channels, and conversion evidence.",
            "Startups: knowledge-graph products for regulated industries where source traceability matters.",
            "Workforce: graph data curator, ontology designer, evidence auditor.",
        ],
        "fresh": "Stress evidence weights: observed, inferred, conceptual, and AI-suggested edges should look different.",
    },
    {
        "slug": "processing_earthquake",
        "title": "First AI Visualization: Earthquake Globe",
        "quick": "The earthquake globe shows how scientific data can become visible and discussable through 3D position, depth, magnitude, time, and sound.",
        "deep": "This is the origin-story topic. I would frame it as a creative prototype that should now be rebuilt as a reproducible data tool: same visual imagination, better provenance.",
        "doe": "DOE relevance: visualization helps technical and nontechnical audiences discuss complex spatial-temporal data before they move into formal modeling.",
        "models": [
            "Poisson / negative-binomial GLM: simple event-count models for region-time windows.",
            "LightGBM anomaly ranker: ranks unusual windows from features such as counts, magnitude bins, depth bins, and clusters.",
            "Hawkes process / ST-DBSCAN: models or groups event sequences that cluster in space and time.",
        ],
        "questions": [
            ("Is this earthquake prediction?", "No. I would explicitly say this is visualization and pattern inspection, not forecasting."),
            ("What makes it scientific instead of just cool?", "Reproducible data pulls, clear encodings, filters, source metadata, and no overclaiming."),
            ("Why sonification?", "Sound can make temporal or magnitude patterns noticeable, but it must be explained and not dramatized."),
        ],
        "sectors": [
            "Finance: market-event visualization without pretending to forecast.",
            "Marketing: live campaign maps showing where activity clusters over time.",
            "Public safety: event dashboards for alerts, inspection, and triage.",
            "Education: interactive science communication tools.",
        ],
        "fresh": "Say the prototype is valuable because it teaches how to turn raw data into intuition, then into a reproducible system.",
    },
    {
        "slug": "seismic",
        "title": "Seismic Notebooks And Pondicherry",
        "quick": "Seismic notebooks can become repeatable workflows for waveform access, arrival picking, QA, maps, and uncertainty summaries.",
        "deep": "I would explain this as notebook-to-system maturity. The original workflow has catalog search, station metadata, waveform downloads, arrival picks, and velocity outputs. AI should help with repetitive QA, not hide uncertainty.",
        "doe": "DOE relevance: DOE-funded subsurface, geothermal, induced-seismicity, carbon storage, and monitoring projects all need scalable waveform processing with provenance.",
        "models": [
            "SeisLM: a seismic foundation model that learns waveform representations from large open waveform collections.",
            "PhaseNet / EQTransformer: neural phase-picking models that propose P and S arrival times.",
            "LightGBM waveform QA: a tree model that flags whether a waveform is usable, weak, mismatched, or needs review.",
        ],
        "questions": [
            ("Where should ML assist first?", "Event triage, station matching, waveform QA, and suggested picks with confidence bands."),
            ("Where should the expert overrule it?", "Weak SNR, bad metadata, unusual paths, and any interpretation that changes the geologic claim."),
            ("How do you show uncertainty?", "Use pick bands, QA colors, velocity residuals, and a reviewed/not-reviewed label."),
        ],
        "sectors": [
            "Healthcare: ECG/EEG signal review with clinician override.",
            "Infrastructure: bridge, rail, and pipeline vibration monitoring.",
            "Finance: noisy time-series anomaly triage with confidence bands.",
            "Workforce: scientific data QA specialist and model-assisted analyst.",
        ],
        "fresh": "The key talking line: model proposes, uncertainty stays visible, human reviews.",
    },
    {
        "slug": "north_slope",
        "title": "AI For Energy Screening Workflows",
        "quick": "Public North Slope geology, maps, wells, logs, papers, and GIS layers can become a source-backed hydrate screening workspace.",
        "deep": "This is the DOE anchor topic. I would say: the current site is not claiming a production model. It is a scaffold for public-data screening. The hard part is building a feature table an expert trusts.",
        "doe": "DOE relevance: hydrate screening, methane risk, energy planning, public data reuse, and subsurface decision support all depend on traceable sources and uncertainty-aware ranking.",
        "models": [
            "Ridge / ElasticNet baseline: simple regression first, so complex models must prove they add value.",
            "Keras ANN hydrate model: a neural network that learns nonlinear combinations of well-log curves to estimate Sgh or screen intervals.",
            "XGBoost / LightGBM challenger: tree-ensemble models for structured well-log tables.",
            "Leave-well-out validation: train on some wells and test on a different well to reduce memorization.",
        ],
        "questions": [
            ("How far can public data go?", "Far enough for screening, source organization, and hypothesis ranking, but not enough for investment-grade resource assessment."),
            ("What makes a hydrate ranking trustworthy?", "Provenance, depth alignment, log QC, held-out wells, uncertainty, and expert review."),
            ("Where is the value?", "Faster screening, fewer bad targets, better source traceability, and better communication between scientists and decision-makers."),
        ],
        "sectors": [
            "Finance: deal screening from filings, reports, market data, and risk signals.",
            "Marketing: segment ranking from CRM, web, social, and campaign data.",
            "AI startups: niche public-data workbenches for regulated sectors.",
            "Jobs: provenance analyst, energy data curator, subsurface ML product manager.",
        ],
        "fresh": "Say provenance is the product. A ranked interval is only useful if someone can ask: where did each feature come from?",
    },
    {
        "slug": "rock_classification",
        "title": "AI For Visual Geoscience Classification",
        "quick": "Rock classification needs images, chemistry, formation context, maps, and expert labels stored together.",
        "deep": "I would frame this as multimodal geoscience. A rock label is not just an image classification problem. It can include thin section texture, geochemistry, spider diagrams, mapped unit, and source notes.",
        "doe": "DOE relevance: resource mapping, critical-mineral exploration support, geologic data integration, and training scientific datasets all need trustworthy labels.",
        "models": [
            "EfficientNet / ResNet: image models that learn visual texture and shape features.",
            "XGBoost / LightGBM chemistry branch: tree models for tabular chemistry and formation variables.",
            "Late-fusion MLP: combines image, chemistry, and text branches after each branch is checked.",
            "CLIP / SigLIP retrieval: finds similar expert-labeled examples instead of asserting one label.",
        ],
        "questions": [
            ("What labels would make this a real dataset?", "Sample ID, scale, source, expert label, ambiguous class flag, chemistry, and formation context."),
            ("Can instruments become rock-property maps?", "They can suggest candidate zones, but labels and geology need expert validation."),
            ("What is the risk?", "Same-sample leakage, weak labels, missing scale, and overconfident mineral claims."),
        ],
        "sectors": [
            "Manufacturing: defect classification from images plus sensor data.",
            "Marketing: customer type classification from images, text, behavior, and purchases.",
            "Finance: company classification from filings, charts, news, and metrics.",
            "Education: AI lab assistant or geology tutor with reviewed examples.",
        ],
        "fresh": "Use the phrase multimodal evidence, not just image recognition.",
    },
    {
        "slug": "valles",
        "title": "SAGE / Valles Caldera Geophysics",
        "quick": "Field geophysics compares imperfect methods: gravity, EM, ERT, TEM, seismic, maps, terrain, and notes.",
        "deep": "This topic should be about disagreement, not just fusion. AI should show where methods agree, where they conflict, and where the evidence is unresolved.",
        "doe": "DOE relevance: geothermal, groundwater, critical minerals, and subsurface energy projects all depend on multi-method interpretation without false certainty.",
        "models": [
            "LightGBM conflict ranker: ranks zones as agreement, conflict, insufficient evidence, or needs review.",
            "Gaussian Process uncertainty: estimates a surface and where the surface is uncertain.",
            "U-Net segmentation: possible future tool for outlining anomaly zones when labels exist.",
            "Survey graph model: represents survey lines and overlaps as nodes and edges.",
        ],
        "questions": [
            ("When does data fusion mislead?", "When it averages away method physics, acquisition limits, or real disagreement."),
            ("What would a field geophysicist want first?", "Method-specific layers, uncertainty, registration checks, and conflict zones."),
            ("What output is honest?", "A review-priority board, not one final clean subsurface truth."),
        ],
        "sectors": [
            "Healthcare: combining imaging, lab tests, history, and notes while keeping conflicts visible.",
            "Finance: conflicting signals from fundamentals, technicals, sentiment, and macro data.",
            "Urban planning: layered climate, traffic, infrastructure, and demographic risk maps.",
            "Workforce: model-risk reviewer and multi-source data fusion analyst.",
        ],
        "fresh": "Strong phrase: unresolved zones should stay unresolved.",
    },
    {
        "slug": "near_surface",
        "title": "AI For Near-Surface Geophysics",
        "quick": "Near-surface field surveys need line-aware comparison across hammer seismic, ERT, TEM, mapped units, and field notes.",
        "deep": "This is the shallow-subsurface version of the Valles issue. The line intersections matter because they create places where methods can check or contradict each other.",
        "doe": "DOE relevance: shallow geophysics supports water, environmental monitoring, geothermal site context, and field-scale method validation.",
        "models": [
            "Line geometry features: line ID, intersection ID, depth, method type, and field-note completeness become model inputs.",
            "LightGBM method-conflict classifier: ranks intervals as agreement, conflict, missing context, or review target.",
            "Gaussian Process / Bayesian surface: method-specific uncertainty surface.",
            "Leave-line-out validation: train on some survey lines and test on another.",
        ],
        "questions": [
            ("Which method should lead?", "It depends on the target property. The site should keep velocity, resistivity, conductivity, and field notes separate first."),
            ("How do line intersections help?", "They are natural validation points where independent methods meet."),
            ("Where should AI leave geology unresolved?", "Where methods conflict, geometry is uncertain, or field notes are missing."),
        ],
        "sectors": [
            "Construction: subsurface risk screening before drilling or building.",
            "Insurance: layered hazard and infrastructure condition maps.",
            "Agriculture/water: shallow groundwater and soil-zone monitoring.",
            "Jobs: field data integrator and uncertainty mapping analyst.",
        ],
        "fresh": "Use line-aware validation as the memorable concept.",
    },
    {
        "slug": "moho_ml",
        "title": "Supervised ML For Moho Depth Mapping",
        "quick": "A model trained on Australia gravity/Moho relationships must prove it transfers before being trusted on the USA.",
        "deep": "This is the cleanest example of real supervised ML. The important discussion is not just what model was used. It is whether the model learned geology or memorized geography.",
        "doe": "DOE relevance: subsurface and energy models often need regional transfer. Honest validation prevents impressive but brittle ML.",
        "models": [
            "Ridge / GAM reference: transparent baseline before complex ML.",
            "LightGBM / XGBoost regressor: tree ensembles for continuous Moho depth prediction.",
            "Keras ANN: neural network challenger after split design is credible.",
            "Gaussian Process residuals: maps where the model misses and where uncertainty is systematic.",
        ],
        "questions": [
            ("How do you test transfer?", "Hold out geography, not random points. Then map residuals by region."),
            ("Why try a simple model first?", "If a complex model cannot beat a transparent reference, the complexity is not justified."),
            ("What matters more than R-squared?", "Residual maps, spatial leakage checks, variable consistency, and geologic plausibility."),
        ],
        "sectors": [
            "Finance: models that work in one market regime but fail in another.",
            "Marketing: customer models transferred to a new region or channel.",
            "Manufacturing: model transfer from one plant or sensor setup to another.",
            "Workforce: ML validation specialist focused on transfer and leakage.",
        ],
        "fresh": "Say: high score is not proof of transfer.",
    },
    {
        "slug": "ambient_noise",
        "title": "AI For Ambient-Noise Seismology",
        "quick": "Ambient-noise processing creates many station pairs and stacks; AI can help triage which correlations are stable enough to review.",
        "deep": "This topic is about scale. Instead of one waveform, ambient noise generates many files, station pairs, cross-correlations, stacks, and parameter choices. AI can help manage QC and monitoring.",
        "doe": "DOE relevance: continuous monitoring matters for geothermal, carbon storage, induced seismicity, and infrastructure-style subsurface observation.",
        "models": [
            "SeisLM-style embedding: pretrained waveform features for quality or change detection.",
            "LightGBM CCF-quality classifier: scores station-pair cross-correlations using stack count, peak strength, symmetry, and metadata.",
            "Isolation Forest: anomaly triage for unusual station-pair windows.",
            "Freshness / seasonal gate: blocks claims from stale data or seasonal/instrument effects.",
        ],
        "questions": [
            ("Which station pairs should be trusted first?", "High stack count, stable peak, good metadata, repeatability, and no obvious instrument issue."),
            ("Where can AI help compute cost?", "Prioritize which pairs, windows, or parameters deserve processing or reprocessing."),
            ("What should stay visible?", "Processing parameters, data freshness, station metadata, and QC flags."),
        ],
        "sectors": [
            "Cybersecurity: continuous log monitoring with anomaly triage.",
            "Manufacturing: sensor-pair monitoring and predictive maintenance.",
            "Finance: monitoring market microstructure or portfolio drift.",
            "Jobs: pipeline reliability analyst and monitoring ML operator.",
        ],
        "fresh": "Differentiate this from the Pondicherry topic: this is continuous workflow management at scale.",
    },
    {
        "slug": "stock_workflow",
        "title": "AI App Building, Automation, And Model Risk",
        "quick": "Codex can turn local notebooks and files into a tracked app, but the app must expose leakage, baselines, and walk-forward validation.",
        "deep": "This is not a finance recommendation. It is a software and model-risk case study. The point is that young builders can ship tools faster, but must learn evidence discipline.",
        "doe": "DOE relevance: the same app-building pattern applies to internal dashboards, monitoring tools, model cards, and scientific workflow automation.",
        "models": [
            "Persistence baseline: simple recent-past rule that fancy models must beat.",
            "ElasticNet: linear model that shrinks weak signals.",
            "LightGBM challenger: nonlinear tree model after the baseline is established.",
            "Walk-forward validation: train on the past, test on the next window, move forward, repeat.",
        ],
        "questions": [
            ("How do we keep model testing honest?", "Past-only features, baseline first, walk-forward validation, drift checks, and cautious language."),
            ("Why include this in a DOE talk?", "It shows the transferable app-building workflow: files -> pipeline -> dashboard -> validation gates."),
            ("What belongs in GitHub/cloud?", "Repeatable scripts, refresh logic, environment notes, tests, and data provenance."),
        ],
        "sectors": [
            "Marketing: campaign dashboards with holdout testing and drift checks.",
            "Commerce: demand forecasting tools with leakage gates.",
            "Startups: rapid vertical SaaS prototypes with model-risk controls.",
            "Jobs: AI app builder, model-risk analyst, data product operator.",
        ],
        "fresh": "Use this as the workforce-development topic: speed is real, but evaluation habits matter.",
    },
    {
        "slug": "sem_petrography",
        "title": "AI For SEM Petrography And Climate Proxies",
        "quick": "SEM image models can propose visible mineral or texture labels, but interpretation and climate-proxy claims need expert and literature support.",
        "deep": "This topic is about separating observation from interpretation. A model might label texture or morphology, but it should not jump straight to paleoclimate or reservoir claims.",
        "doe": "DOE relevance: materials, subsurface, environmental, and reservoir workflows all need image-based labels with traceable interpretation boundaries.",
        "models": [
            "Image classifier: proposes visible labels from SEM crops.",
            "Segmentation model: outlines grains, pores, textures, or mineral patches when masks exist.",
            "Image-text retrieval: finds similar expert-labeled examples and literature-linked descriptions.",
        ],
        "questions": [
            ("What labels are defensible?", "Visible labels: texture, morphology, mineral candidate, scale, and sample context."),
            ("Where does interpretation begin?", "When the label becomes a climate, reservoir, or process claim, evidence requirements increase."),
            ("What examples would the model need?", "Expert-reviewed crops, scale bars, sample metadata, literature-linked labels, and ambiguous examples."),
        ],
        "sectors": [
            "Materials science: microscopy defect and phase labeling.",
            "Healthcare: pathology-style image review with expert confirmation.",
            "Manufacturing: microscope QA for product defects.",
            "Jobs: image-label auditor and scientific dataset curator.",
        ],
        "fresh": "Memorable phrase: visible label first, proxy claim later.",
    },
]


FULL_TOUR_SCRIPT = [
    "Thank you for looking through this portfolio with me. I am going to move quickly across the whole website first, because the projects are different on the surface, but they are really connected by one workflow: take messy scientific artifacts, organize them into evidence, use AI or machine learning where it helps, and keep the human review step visible.",
    "The first topic, AI agents learning scientific software, is about how tools like Codex could eventually learn from expert demonstrations. In a DOE context, that matters because scientific work often happens inside specialized software, not just in clean text prompts. A useful agent would need to see files, screen states, map layers, errors, and final outputs, then be scored against a rubric.",
    "The critical-minerals knowledge graph topic is about turning papers, deposits, minerals, host rocks, and evidence into a connected map. I would not present this as AI magically discovering minerals. I would present it as a reviewable way to trace relationships and ask better questions about supply chains and energy-relevant materials.",
    "The earthquake globe is the creative origin point. It shows how spatial-temporal data can become understandable through position, color, magnitude, depth, time, and even sound. The important point is not earthquake prediction. The important point is moving from an exciting visualization toward a reproducible scientific communication tool.",
    "The seismic notebooks topic shows how exploratory geophysics can become a repeatable workflow: catalog search, waveform download, arrival picking, quality control, velocity interpretation, and uncertainty. Machine learning can help with triage and phase picking, but the expert still decides what is geologically defensible.",
    "For North Slope energy screening, the website starts to look very DOE-relevant. The question is whether public geology, maps, source libraries, and ML planning can become a screening atlas for hydrate or subsurface energy intervals. The model is not the whole answer. The value is in the feature table, the validation plan, and showing uncertainty before anyone makes a decision.",
    "The rock classification topic connects petrography, geochemistry, diagrams, and labeled examples. I would explain this as a multimodal classification problem: images, chemistry, texture, and expert labels need to be stored together before any model becomes useful.",
    "The Valles Caldera and near-surface topics are about comparing imperfect geophysical methods. Gravity, EM, ERT, seismic, field notes, and geologic maps do not always agree. AI can help rank conflicts, align lines, and flag uncertainty, but it should not erase disagreement just to make a clean-looking map.",
    "The Moho depth mapping topic is the clearest machine-learning validation story. Training on Australia and testing transfer to the United States asks whether a model learned a real geophysical relationship or just memorized a region. That is a very useful point for DOE, because energy and subsurface models often have to move across regions where data quality changes.",
    "The ambient-noise topic is about scale. Continuous seismic records generate huge numbers of station pairs, cross-correlations, stacks, and monitoring outputs. AI can help decide which correlations are trustworthy enough to review first, but the pipeline must keep compute choices, station metadata, and QC flags visible.",
    "The stock workflow is included because it shows the workforce side of AI: using Codex to build real apps quickly while still respecting baselines, leakage, validation, and model risk. Even outside geoscience, the same discipline applies: build fast, test honestly, and do not let a polished dashboard hide weak assumptions.",
    "The SEM petrography topic is about image-based scientific labeling. A model might label visible textures or mineral candidates, but it should not jump straight to climate or reservoir claims. I would say: visible label first, proxy claim later.",
    "So if I summarize the whole website in one sentence, it is this: AI should help move scientific work from scattered artifacts into reviewable systems, where evidence, models, uncertainty, and human judgment are all visible.",
]


TOPIC_READALOUD = {
    "ai_workflow": [
        "For this topic, I would say: this is about training AI agents on the actual work scientists do in software. A lot of science is not just text. It is opening files, checking coordinate systems, moving between notebooks and maps, exporting figures, and deciding whether the output is scientifically acceptable.",
        "The model idea here is behavior cloning. Instead of only asking a model a question, we record expert demonstrations: what was on the screen, what file was open, what action happened next, and what counted as success. The DOE angle is that scientific software workflows are expensive to teach and easy to break, so the evaluation rubric matters as much as the agent.",
    ],
    "thesis_graph": [
        "For the critical-minerals graph, I would frame it like this: papers and reports contain valuable relationships, but those relationships are locked inside prose. A graph turns them into a structure: deposit, mineral, host rock, alteration, source, confidence, and evidence.",
        "The model is not replacing a geologist. It is helping extract entities, rank relationships, and connect every claim back to a source. That is why GraphRAG is useful here: it gives an answer by walking through evidence links instead of just producing a fluent paragraph.",
    ],
    "processing_earthquake": [
        "For the earthquake globe, I would say this is my first example of turning data into intuition. Earthquakes have location, depth, magnitude, and time, and a 3D globe makes those variables easier to discuss than a table.",
        "The careful part is the boundary. This is not prediction. If I add machine learning, it would be for clustering, anomaly ranking, or summarizing patterns for review. The value is reproducible communication: same data pull, same encoding rules, clear source, and no overclaiming.",
    ],
    "seismic": [
        "For the seismic notebooks, I would explain that the workflow starts as exploration but should mature into a pipeline. You search a catalog, download waveforms, inspect quality, pick arrivals, calculate outputs, and document uncertainty.",
        "The ML models here, like PhaseNet or EQTransformer, can propose arrival picks. But I would tell the audience that the model is an assistant. It can speed up repetitive picks and flag weak waveforms, while the expert still handles ambiguous signals and geologic interpretation.",
    ],
    "north_slope": [
        "For the North Slope section, I would slow down because this is one of the strongest DOE connections. The question is whether public geology, maps, literature, stratigraphy, and ML planning can become a screening workflow for hydrate or subsurface energy targets.",
        "The model I would describe first is a baseline, like Ridge or ElasticNet, because it is easier to explain and audit. Then I would compare it to XGBoost or a neural network. The key is leave-well-out validation: can the model make a reasonable estimate on wells it did not see, instead of memorizing a known location?",
    ],
    "rock_classification": [
        "For rock classification, I would say the website is showing the raw ingredients of a future model: thin sections, chemical diagrams, formation tables, and labels. The important step is not throwing all of that into AI immediately. It is defining what the label means.",
        "A good workflow could combine an image model for visual texture, a tree model for chemistry, and a fusion model that uses both. I would explain that multimodal AI only works if the examples are aligned: the image, the chemistry, the location, the expert label, and the uncertainty all need to refer to the same sample.",
    ],
    "valles": [
        "For Valles Caldera, I would say this is about field geophysics and interpretation under uncertainty. Different methods see different parts of the subsurface, and they do not always agree.",
        "The model concept is data fusion with uncertainty. A U-Net might segment anomalies on a map, a Gaussian Process might show where uncertainty is high, and a conflict ranker might tell the interpreter which locations deserve attention first. The goal is not a prettier map; the goal is a more honest comparison.",
    ],
    "near_surface": [
        "For near-surface geophysics, I would make the same idea more concrete. Hammer seismic, ERT, TEM, geologic units, and line intersections all provide partial views. The question is where they agree and where they contradict each other.",
        "A useful model would not just predict a layer. It would flag line intersections that do not line up, show method limits, and identify which areas need more field review. That is valuable because shallow decisions can affect environmental work, infrastructure, geothermal siting, and water questions.",
    ],
    "moho_ml": [
        "For Moho ML, I would say this is my cleanest example of honest machine-learning testing. Training on Australia and testing on the United States asks a hard question: did the model learn physics-related structure, or did it learn regional shortcuts?",
        "That is why I would start with a simple model, then compare it to LightGBM or a neural network. The residual map becomes a scientific tool. Where the model is wrong, the error is not just failure; it may point to geology, missing features, or transfer limits.",
    ],
    "ambient_noise": [
        "For ambient-noise seismology, I would explain the scale problem. Continuous station data becomes station pairs, cross-correlations, stacks, monitoring windows, and quality-control decisions. The number of files can explode before the scientist even knows what is trustworthy.",
        "The ML idea is a quality gate. A model can rank which cross-correlations look stable, which station pairs are suspicious, and which changes might be seasonal or instrument-related. The expert still decides what becomes a scientific interpretation.",
    ],
    "stock_workflow": [
        "For the stock workflow, I would say this is not here because stock prediction is the same as geoscience. It is here because the workflow discipline transfers: build an app, define features, set a baseline, avoid leakage, validate over time, and show risk.",
        "This is also a jobs topic. The new skill is not only coding a model; it is building a usable tool with version control, data provenance, tests, and a clear explanation of what the model can and cannot support.",
    ],
    "sem_petrography": [
        "For SEM petrography, I would say the main discipline is separating observation from interpretation. A model can help label visible textures, pores, grains, or mineral candidates, but climate-proxy or reservoir conclusions require more evidence.",
        "The useful workflow is expert-reviewed image crops, scale bars, sample metadata, and literature-linked definitions. My phrase for this slide would be: visible label first, proxy claim later.",
    ],
}


SPECIFIC_PIVOTS = {
    "ai_workflow": [
        "Finance: an agent watches a senior analyst build a credit memo in Excel, Bloomberg, SEC filings, and PowerPoint, then later recreates the packet with every source and formula logged.",
        "Marketing operations: an agent learns the exact workflow for pulling Meta, Google Ads, Shopify, and HubSpot data into a weekly campaign report, then flags missing UTMs or broken dashboards before a manager reviews it.",
        "Healthcare administration: the same screen-and-rubric approach could train agents to prepare prior-authorization packets, with humans approving anything patient-specific.",
        "Startup opportunity: vertical AI agents for expensive niche software such as GIS, CAD, lab instruments, and compliance tools, sold with audit trails instead of generic chat.",
    ],
    "thesis_graph": [
        "Finance: a supply-chain risk graph connecting mining companies, refineries, customers, tariffs, shipping chokepoints, and SEC risk disclosures for battery or rare-earth exposure.",
        "Marketing: a customer evidence graph linking persona, pain point, campaign touch, product feature, sales objection, and conversion outcome so teams see why messages work.",
        "AI startup ecosystem: graph products for regulated markets where every generated answer must cite the relationship and the original document, such as energy permitting or insurance.",
        "Jobs: ontology designer, evidence librarian, graph QA analyst, and domain data curator become practical roles because the schema determines whether AI output is trustworthy.",
    ],
    "processing_earthquake": [
        "Finance: event-globe logic becomes a market-shock room showing where earnings misses, rate changes, oil shocks, and volatility clusters happened over time without claiming prediction.",
        "Retail and commerce: a live order-delay map showing where warehouse delays, weather, and carrier exceptions cluster so teams can triage customer messaging.",
        "Marketing: launch monitoring that visualizes attention spikes by city, channel, influencer, and time, then separates real demand from noisy one-time bursts.",
        "Public-sector operations: emergency dashboards for inspections, infrastructure calls, or outage reports where visual triage matters before formal modeling.",
    ],
    "seismic": [
        "Energy operations: a waveform-QA pipeline for geothermal or carbon-storage monitoring where weak signals are ranked before expert review.",
        "Finance/compliance analogy: transaction-monitoring alerts use the same pattern: rank noisy events, show confidence, preserve the human decision, and audit the review.",
        "Manufacturing: vibration sensors on machines can use similar signal models to flag bearing failures or abnormal cycles before downtime.",
        "Jobs: signal-QA analyst and scientific pipeline operator, combining domain judgment with model monitoring.",
    ],
    "north_slope": [
        "Energy finance: lenders or investors could use a transparent screening atlas to compare acreage, public geology, infrastructure distance, and uncertainty before funding diligence.",
        "Insurance: subsurface risk scoring for pipelines, permafrost exposure, or infrastructure planning, using maps plus uncertainty instead of a single black-box score.",
        "AI startup: a commercial due-diligence product that turns public geology and documents into first-pass screening dashboards for geothermal, minerals, hydrogen storage, or carbon storage.",
        "Workforce: energy data product manager, geoscience ML validator, and public-data integration specialist are realistic job lanes.",
    ],
    "rock_classification": [
        "Mining and materials: ore sorting or core logging workflows where images, chemistry, and location are fused to prioritize samples for expert review.",
        "Construction commerce: aggregate suppliers could classify material quality from photos plus lab chemistry before sending batches to infrastructure projects.",
        "Marketing/content: product teams selling lab instruments could use labeled visual examples to explain which measurements matter to non-experts.",
        "Jobs: multimodal dataset builder, lab-data label manager, and geoscience annotation lead.",
    ],
    "valles": [
        "Geothermal development: ranking areas where gravity, resistivity, and geologic maps agree or conflict before drilling money is spent.",
        "Infrastructure: city planners comparing shallow hazard maps, utility maps, and geotechnical surveys before road, bridge, or tunnel work.",
        "Insurance and risk: uncertainty-aware maps for wildfire recovery, landslide exposure, or subsidence risk where a clean map can be misleading.",
        "Startup: field-data fusion software that sells the disagreement view, not just the final interpreted map.",
    ],
    "near_surface": [
        "Environmental consulting: comparing ERT, TEM, soil borings, and field notes for contamination plumes or wetland boundaries.",
        "Real estate and infrastructure: shallow subsurface screening before solar farms, data centers, highways, or water projects.",
        "Agriculture and water: mapping shallow moisture or salinity patterns by fusing sensor lines, soil maps, and field observations.",
        "Jobs: near-surface data integrator, field QA technologist, and uncertainty map reviewer.",
    ],
    "moho_ml": [
        "Finance: model-transfer testing is like training a credit model in one economic cycle and testing it in another; the transfer split is the real honesty check.",
        "Marketing: a conversion model trained on one region or audience must be tested on a new region before a company trusts the budget recommendation.",
        "Energy exploration: models trained in one basin need leave-region-out validation before being used to screen a new basin.",
        "AI startups: selling model-validation services that produce residual maps, failure modes, and transfer-risk reports instead of only accuracy scores.",
    ],
    "ambient_noise": [
        "Industrial IoT: continuous sensor streams from factories can use similar quality gates to decide which machine signals deserve human maintenance review.",
        "Cybersecurity: log streams become event correlations; models rank which correlations are stable threats versus routine or seasonal noise.",
        "E-commerce operations: continuous traffic, checkout errors, inventory changes, and ad spend can be monitored with anomaly gates before managers react.",
        "Jobs: monitoring-ML operator, alert-quality analyst, and model-drift reviewer.",
    ],
    "stock_workflow": [
        "Marketing analytics: a campaign forecast app should compare against a simple baseline, use holdout periods, and warn when new channels make old patterns unreliable.",
        "Retail commerce: demand forecasting for inventory should include leakage checks so future promotions or future stockouts do not accidentally enter training data.",
        "AI startup operations: founders can use Codex-style workflows to build internal tools quickly, but investor-ready products need testing, logging, and repeatable deployment.",
        "Jobs: AI product operator, validation-focused app builder, model-risk analyst, and GitHub-literate business analyst.",
    ],
    "sem_petrography": [
        "Materials manufacturing: SEM defect detection for battery materials, semiconductors, coatings, or catalysts, with expert confirmation before process changes.",
        "Healthcare analogy: pathology image triage has the same boundary: model labels visible patterns, but diagnosis remains expert-reviewed.",
        "Carbon and climate tech: mineral texture labeling can support weathering or sequestration research if the proxy claim links back to sample context and literature.",
        "Jobs: microscopy data curator, image-label QA specialist, scientific annotation lead, and lab-AI workflow designer.",
    ],
}


def configure_document(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style in ["List Bullet", "List Number"]:
        style = styles[list_style]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("DOE AI Portfolio Presenter Guide")
    style_run(run, color=MUTED, size=9)


def add_title(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("DOE AI Portfolio Presenter Guide")
    style_run(r, bold=True, color=INK, size=24)
    p2 = doc.add_paragraph()
    r2 = p2.add_run("Fast tour script, deep-dive topic notes, Q&A pivots, and cross-sector applications")
    style_run(r2, color=MUTED, size=12)
    add_callout(
        doc,
        "Core thesis",
        "AI is most useful here when it turns scattered scientific evidence into reviewable workflows: source -> feature -> model -> uncertainty -> human decision.",
        fill=LIGHT_BLUE,
    )


def add_quick_pass(doc):
    add_heading(doc, "Fast Pass-Through Script", 1)
    add_body(
        doc,
        "Use this if you need to show every topic quickly. Spend about 20 to 30 seconds per topic, then invite the audience to pick where to go deeper.",
    )
    add_body(
        doc,
        "Opening: Thank you for taking a look at this. The portfolio is not meant to claim that every project is finished. It is meant to show how AI can help organize scientific work into systems that preserve evidence, uncertainty, and expert review.",
    )
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [2500, 4560, 2300])
    hdr = table.rows[0].cells
    for i, text in enumerate(["Topic", "30-second script", "Natural question"]):
        set_cell_shading(hdr[i], LIGHT_BLUE)
        hdr[i].paragraphs[0].add_run(text).bold = True
    for topic in TOPICS:
        row = table.add_row().cells
        row[0].text = topic["title"]
        row[1].text = topic["quick"]
        row[2].text = topic["questions"][0][0]
    doc.add_paragraph()
    add_callout(
        doc,
        "Fast transition line",
        "Across all of these, I am asking the same question: what evidence enters the system, what model or workflow acts on it, what uncertainty remains, and where does the human expert step in?",
        fill=LIGHT_GRAY,
    )


def add_full_tour_script(doc):
    add_heading(doc, "10-15 Minute Read-Aloud Website Tour", 1)
    add_callout(
        doc,
        "How to use this",
        "Read this straight through if you need to cover the whole website in one pass. It is written for about 10 to 15 minutes depending on how much you pause on the visuals.",
        fill=LIGHT_BLUE,
    )
    for paragraph in FULL_TOUR_SCRIPT:
        add_number(doc, paragraph)
    add_callout(
        doc,
        "Bridge into questions",
        "After the tour, say: I can go deeper on any one of these, but the main thing I want you to notice is that the same AI pattern keeps repeating: evidence in, model or workflow in the middle, uncertainty out, and human review before decisions.",
        fill=LIGHT_GRAY,
    )


def add_deep_order(doc):
    add_heading(doc, "Recommended DOE Deep-Dive Order", 1)
    add_body(doc, "If time is limited, use this order because it moves from DOE-relevant energy data into scientific ML validation and then workforce implications.")
    for item in [
        "North Slope energy screening: public data, hydrate screening, provenance, and expert review.",
        "Valles / near-surface geophysics: method comparison, disagreement, and uncertainty-aware field interpretation.",
        "Seismic / ambient noise: scalable processing, waveform QA, and monitoring workflows.",
        "Knowledge graphs for critical minerals: evidence-backed relationships and source traceability.",
        "Moho ML / stock workflow: honest transfer testing, leakage, baselines, and model-risk culture.",
    ]:
        add_number(doc, item)


def add_topic_section(doc, topic):
    add_heading(doc, topic["title"], 1)
    add_callout(doc, "One-sentence anchor", topic["quick"], fill=LIGHT_BLUE)

    add_heading(doc, "Read-Aloud Topic Script", 2)
    for paragraph in TOPIC_READALOUD.get(topic["slug"], [topic["deep"]]):
        add_body(doc, paragraph)

    add_heading(doc, "Backup Speaker Notes", 2)
    add_body(doc, topic["deep"])
    add_body(doc, topic["doe"], bold_prefix="DOE relevance:")
    add_body(doc, "Fresh angle: " + topic["fresh"], bold_prefix="Fresh angle:")

    add_heading(doc, "Model Terms To Explain Out Loud", 2)
    for model in topic["models"]:
        add_bullet(doc, model)

    add_heading(doc, "If Someone Raises Their Hand", 2)
    qa_table = doc.add_table(rows=1, cols=2)
    set_table_geometry(qa_table, [3100, 6260])
    qa_hdr = qa_table.rows[0].cells
    qa_hdr[0].text = "Likely question"
    qa_hdr[1].text = "Answer you can give"
    for cell in qa_hdr:
        set_cell_shading(cell, LIGHT_BLUE)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for question, answer in topic["questions"]:
        row = qa_table.add_row().cells
        row[0].text = question
        row[1].text = answer
    doc.add_paragraph()

    add_heading(doc, "Specific Cross-Sector Pivots", 2)
    for sector in SPECIFIC_PIVOTS.get(topic["slug"], topic["sectors"]):
        add_bullet(doc, sector)

    add_heading(doc, "Safe Claim Boundary", 2)
    add_body(
        doc,
        "Do not overclaim. Say this is a workflow, scaffold, review system, or prototype unless the data, validation split, and expert review support a stronger statement.",
    )


def add_closing(doc):
    add_heading(doc, "Closing Script", 1)
    add_body(
        doc,
        "My closing point is that the model name is not the main story. The main story is the workflow around the model: source evidence, feature design, validation, uncertainty, and expert review.",
    )
    add_body(
        doc,
        "For DOE and energy work, that matters because public data, scientific software, field methods, and subsurface models are only useful if the reasoning remains inspectable. AI can make that work faster, but the human review loop is what makes it trustworthy.",
    )
    add_body(
        doc,
        "If I had to summarize the portfolio in one sentence: AI should help us move from scattered artifacts to reviewable scientific systems, not from messy data to unsupported certainty.",
    )


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_title(doc)
    add_quick_pass(doc)
    add_full_tour_script(doc)
    add_deep_order(doc)
    doc.add_page_break()
    add_heading(doc, "Deep Topic-by-Topic Script", 1)
    add_body(doc, "Use these when a topic gets selected for deeper discussion. Each topic has a speaker script, DOE framing, model translation, likely questions, and cross-sector pivots.")
    for topic in TOPICS:
        add_topic_section(doc, topic)
    add_closing(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
